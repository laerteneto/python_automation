from docx import Document
from docx.shared import Inches


class EvidenceGenerator:
    def __init__(self, project, exec_time, final_result):
        self.document = Document()
        self.FillHeader(project, exec_time, final_result)
        
    def FillHeader(self, project, exec_time, final_result):
        self.document.add_heading('Evidence document', 0)
        records = (
            (project, exec_time, final_result),
        )   
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Project'
        hdr_cells[1].text = 'Execution time'
        hdr_cells[2].text = 'Final Result'
        for qty, _id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = _id
            row_cells[2].text = desc
 
    def AddEvidence(self, test_name, evidence_name, picture):
        self.document.add_heading(test_name, level=1)
        self.document.add_paragraph(evidence_name, style='Intense Quote')
        self.document.add_picture(picture, width=Inches(6.25))
        self.document.add_page_break()

    def CreateDocument(self, doc_test):
        self.document.save(doc_test)
